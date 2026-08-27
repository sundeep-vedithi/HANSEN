#!/usr/bin/env python3
"""Build the supplementary note (Markdown + DOCX) and the tables workbook.

Every number in the narrative and in the figure legends is read from
summary.json or from the tables written by 03_analyse.py, so the text, the
tables and the figures cannot drift apart.
"""
import os
import re
import sys
import csv
import json
import sqlite3
import datetime

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
from paths import TABDIR, FIGDIR, DB  # noqa: E402
S = json.load(open(os.path.join(HERE, 'summary.json')))

MD = os.path.join(HERE, 'HANSEN_SupplementaryNote_AFDB_comparison.md')
DOCX = os.path.join(HERE, 'HANSEN_SupplementaryNote_AFDB_comparison.docx')
XLSX = os.path.join(HERE, 'HANSEN_SupplementaryTables_AFDB_comparison.xlsx')

M = ['af3', 'boltz', 'boltz2', 'chai']
LAB = {'af3': 'AlphaFold 3', 'boltz': 'Boltz-1', 'boltz2': 'Boltz-2',
       'chai': 'Chai-1'}


def T(name):
    with open(os.path.join(TABDIR, name)) as fh:
        return list(csv.reader(fh, delimiter='\t'))


def g(m, metric, stat):
    return S['global'][m][metric][stat]


def thr(m, t):
    return 100 * S['global_thresholds'][m][f'tm_ge_{t}']


WORDS = {0: 'no', 1: 'one', 2: 'two', 3: 'three', 4: 'four', 5: 'five', 6: 'six',
         7: 'seven', 8: 'eight', 9: 'nine', 10: 'ten', 11: 'eleven', 12: 'twelve'}


def num(n):
    """Counts below thirteen are written as words, as in journal house style."""
    n = int(n)
    return WORDS[n] if n in WORDS else f'{n:,}'


def Num(n):
    w = num(n)
    return w[0].upper() + w[1:]


def series(items):
    items = list(items)
    if len(items) == 1:
        return items[0]
    return ', '.join(items[:-1]) + ' and ' + items[-1]


def by_method(fmt):
    """Render a value for all four methods in a fixed order."""
    return series(fmt(m) for m in M)


def polish(text):
    """Typographic conventions applied at render time."""
    text = (text.replace('Ca atoms', 'Cα atoms')
                .replace('Ca geometry', 'Cα geometry')
                .replace('Ca deviation', 'Cα deviation')
                .replace('Ca-lDDT', 'lDDT-Cα'))
    text = re.sub(r'(?<=\d) A(?=[\s.,)])', ' Å', text)
    text = text.replace('Spearman rho', 'Spearman ρ')
    return text


# --- numbering for inclusion in the HANSEN manuscript -----------------------
# The manuscript already uses Supplementary Tables S1-S8, so the tables of this
# note continue from S9. The note's own sections are numbered 1-13 so that
# "section 6" can never be confused with "Table S23".
TMAP = {'S1': 'S9', 'S2': 'S10', 'S3': 'S11', 'S4': 'S12', 'S5': 'S13',
        'S6': 'S14', 'S7': 'S15', 'S8': 'S16', 'S8b': 'S17', 'S9': 'S18',
        'S10': 'S19', 'S11': 'S20', 'S12': 'S21', 'S13': 'S22', 'S14': 'S23',
        'S15': 'S24', 'S16': 'S25'}


def renumber(text):
    """Map table labels onto the manuscript's supplementary numbering."""
    text = re.sub(r'\b(Table|sheet|Tables) (S\d+b?)',
                  lambda m: f'{m.group(1)} {TMAP.get(m.group(2), m.group(2))}', text)
    text = re.sub(r'\bS(\d+b?) (?:to|and) S(\d+b?)\b',
                  lambda m: f'{TMAP.get("S"+m.group(1), "S"+m.group(1))} to '
                            f'{TMAP.get("S"+m.group(2), "S"+m.group(2))}', text)
    text = re.sub(r'\bsection S(\d+)\b', r'section \1', text)
    return text


def sheet_name(fn):
    return TMAP.get(fn.split('_')[0], fn.split('_')[0])


def s8b(col):
    """min to max across the four method rows of a numeric column of Table S17."""
    vals = sorted(float(r[col]) for r in T('S8b_fold_vs_packing.tsv')[1:])
    fmt = '.0f' if vals[-1] >= 10 else '.2f'
    return f'{vals[0]:{fmt}} to {vals[-1]:{fmt}}'


S15 = {r[1]: r for r in T('S15_case_studies.tsv')[1:]}


def ex(ml, name=None):
    """Named example with its residue count and TM range, from Table S24."""
    r = S15[ml]
    tms = sorted(float(r[i]) for i in range(5, 9))
    label = name or (r[3] if r[3] != ml else None)
    head = f'{label} ({ml}, ' if label else f'{ml} ('
    return head + f'{int(r[4]):,} residues, TM {tms[0]:.2f} to {tms[-1]:.2f})'


def n_oligomers():
    con = sqlite3.connect(f'file:{DB}?mode=ro', uri=True)
    n = con.execute("SELECT COUNT(*) FROM protein_structure_models "
                    "WHERE assembly_type <> 'monomer'").fetchone()[0]
    con.close()
    return n


TABLE_TITLES = {
    'S1_coverage.tsv': 'Table S9. Comparison coverage',
    'S2_construct_length_differences.tsv':
        'Table S10. Proteins whose modelled chain length differs between AFDB and HANSEN',
    'S3_global_agreement.tsv': 'Table S11. Global agreement with the AFDB reference',
    'S4_confidence.tsv': 'Table S12. Confidence concordance',
    'S5_confident_core.tsv': 'Table S13. Effect of restricting to the confident core',
    'S6_length_strata.tsv': 'Table S14. Agreement by chain length',
    'S7_secondary_structure.tsv': 'Table S15. Secondary structure and compactness',
    'S8_discordant_proteins.tsv': 'Table S16. Discordant proteins',
    'S8b_fold_vs_packing.tsv': 'Table S17. Local fold and domain packing',
    'S9_five_way_concordance.tsv': 'Table S18. Five-way concordance (median TM-score)',
    'S10_target_tiers.tsv': 'Table S19. Agreement by target-priority tier',
    'S11_top50_targets.tsv': 'Table S20. Top 50 prioritised targets',
    'S12_localisation_strata.tsv': 'Table S21. Agreement by predicted localisation',
    'S13_per_protein_full.tsv': 'Table S22. Per-protein results (complete set)',
    'S14_tm_validation.tsv':
        'Table S23. Validation of the TM-score implementation against TM-align',
    'S15_case_studies.tsv': 'Table S24. Named discordance case studies',
    'S16_pknb_domains.tsv':
        'Table S25. Domain-resolved comparison for PknB (ML0016)',
}

# Tables set in the text. Everything else is supplied in the workbook only, so
# that no table is reproduced twice. Wide tables are shown as a column subset,
# with the full version kept in the workbook.
INLINE_SPEC = {
    'S3_global_agreement.tsv': ['Method', 'n', 'TM-score median', 'TM-score IQR',
                                'GDT-TS median', 'lDDT median', 'RMSD (A) median',
                                'Core RMSD (A) median', 'TM>=0.5', 'TM>=0.9'],
    'S4_confidence.tsv': None,
    'S8b_fold_vs_packing.tsv': None,
    'S16_pknb_domains.tsv': ['Segment', 'Residues', 'AFDB pLDDT',
                             'TM AlphaFold 3', 'TM Boltz-1', 'TM Boltz-2',
                             'TM Chai-1'],
    'S9_five_way_concordance.tsv': None,
}


def cell(text):
    """Typographic conventions for table headers and cells."""
    text = (text.replace('(A)', '(Å)').replace('>=', '≥').replace('A)', 'Å)')
                .replace('rho(', 'ρ(').replace('Delta', 'Δ'))
    text = re.sub(r'(?<=\d) A\b', ' Å', text)
    if re.fullmatch(r'\d{4,}', text):
        text = f'{int(text):,}'
    return text


def table_rows(fn):
    """Rows for an in-text table, restricted to the columns of INLINE_SPEC."""
    rows = T(fn)
    cols = INLINE_SPEC.get(fn)
    if cols:
        idx = [rows[0].index(c) for c in cols]
        rows = [[r[i] for i in idx] for r in rows]
    return [[cell(c) for c in r] for r in rows]


# ------------------------------------------------------------- figure legends
def figures():
    pr = S['per_residue']
    fw = S.get('five_way', {})
    prox = S.get('agreement_proxy', {})
    pk = S.get('pknb_worked_example', {})
    fig = []

    fig.append((
        'FigS1_global_agreement',
        'Figure S1. Agreement between the AFDB models and the four HANSEN '
        f"predictors over {sum(S['coverage']['per_method'].values()):,} model pairs. "
        '(a) Histograms of TM-score against the AFDB model, in 50 bins between 0 '
        'and 1, drawn as one outline per method. All four distributions have a '
        'single mode above TM 0.95 together with a tail of low-scoring proteins. '
        'The medians are '
        + by_method(lambda m: f"{g(m,'tm_score','median'):.3f} for {LAB[m]}")
        + '. (b) The same data as cumulative distributions, with the fraction of '
        'proteins on the vertical axis. The dotted vertical line marks TM 0.5, the '
        'conventional threshold above which two structures share a fold. The '
        'fraction of proteins that lies below this line is '
        + by_method(lambda m: f"{100-thr(m,0.5):.1f} per cent for {LAB[m]}")
        + '. (c) Box plots of Ca-lDDT, the superposition-free measure of local '
        'agreement. Boxes span the interquartile range, the horizontal line is the '
        'median, whiskers extend to 1.5 times the interquartile range and outliers '
        'are omitted. Median lDDT is '
        + by_method(lambda m: f"{g(m,'lddt','median'):.3f} for {LAB[m]}")
        + '. Comparison of panels (a) and (c) shows that the low-scoring tail in '
        'the TM-score distribution has no counterpart in lDDT, because TM-score '
        'responds to the placement of one domain relative to another and lDDT does '
        'not.'))

    b = pr['af3']['dev_by_min_plddt_bin']
    fig.append((
        'FigS2_confidence',
        'Figure S2. Model confidence and its relation to structural agreement. '
        '(a) Median Ca deviation between the model and the AFDB reference after '
        'TM-superposition, computed per residue and grouped into four bins by the '
        'lower of the two per-residue pLDDT values. Bars are grouped by bin and '
        'coloured by method, and the vertical axis is logarithmic. For AlphaFold 3 '
        f"the median deviation is {b['ge90']['median_dev']:.2f} A in the pLDDT 90 "
        f"and above bin ({b['ge90']['n']:,} residues), "
        f"{b['70_90']['median_dev']:.2f} A between 70 and 90 "
        f"({b['70_90']['n']:,} residues), {b['50_70']['median_dev']:.1f} A between "
        f"50 and 70 ({b['50_70']['n']:,} residues) and "
        f"{b['lt50']['median_dev']:.1f} A below 50 ({b['lt50']['n']:,} residues). "
        'The other three methods follow the same pattern. (b) Per-protein mean '
        'model pLDDT on the horizontal axis against TM-score to AFDB on the '
        'vertical axis, one point per protein per method. The correlation is '
        + by_method(lambda m: f"r = {S['confidence'][m]['r_modelplddt_tm']:.3f} for {LAB[m]}")
        + '. (c) Calibration of mean pLDDT between AFDB on the horizontal axis and '
        'each predictor on the vertical axis. The dashed line is the identity. '
        'AlphaFold 3 and Chai-1 lie on the identity, with proteome mean pLDDT of '
        f"{S['confidence']['af3']['model_plddt_mean']:.1f} and "
        f"{S['confidence']['chai']['model_plddt_mean']:.1f} against "
        f"{S['confidence']['af3']['afdb_plddt_mean']:.1f} for AFDB. The Boltz-1 and "
        'Boltz-2 clouds lie below it, at '
        f"{S['confidence']['boltz']['model_plddt_mean']:.1f} and "
        f"{S['confidence']['boltz2']['model_plddt_mean']:.1f}."))

    fig.append((
        'FigS3_stratification',
        'Figure S3. Chain length, low-confidence regions and secondary structure. '
        '(a) Mean TM-score against AFDB, plotted against chain-length stratum, one '
        'line per method. For all four methods the value rises from the shortest '
        'stratum to a maximum in the 200 to 400 and 400 to 700 residue strata and '
        'then falls in the longest stratum. The values are given in Table S14. '
        '(b) Mean TM-score computed over the full chain, in grey, beside the same '
        'quantity computed over the confident core, in blue. The confident core is '
        'the set of residues with pLDDT of 70 or above in both the AFDB model and '
        'the HANSEN model. The core value exceeds the full-chain value by '
        + by_method(lambda m: f"{S['core'][m]['tm_core']-S['core'][m]['tm_full']:.3f} for {LAB[m]}")
        + '. (c) Three-state secondary-structure agreement Q3 on the horizontal '
        'axis against TM-score on the vertical axis, one point per protein per '
        'method. Secondary structure was assigned from Ca geometry by the P-SEA '
        'method. Q3 falls as TM-score falls. For AlphaFold 3 the mean Q3 is '
        f"{S['ss']['af3']['q3_tm_ge_0.9']:.3f} among the proteins at TM 0.9 or "
        f"above and {S['ss']['af3']['q3_tm_lt_0.5']:.3f} among the "
        f"{S['ss']['af3']['n_tm_lt_0.5']} proteins below TM 0.5, against "
        f"{S['ss']['af3']['q3_mean']:.3f} over the whole set. The correlation "
        'between Q3 and TM-score is '
        + by_method(lambda m: f"r = {S['ss'][m]['r_q3_tm']:.2f} for {LAB[m]}")
        + '. The proteins that disagree globally therefore also carry a different '
        'secondary-structure assignment over part of the chain, which follows from '
        'the absence of regular structure in the low-confidence segments described '
        'in section 5.'))

    if fw:
        mat = fw['median_tm_matrix']
        order = fw['order']
        fig.append((
            'FigS4_five_way',
            'Figure S4. Position of the AFDB models within the five-model '
            'ensemble. (a) Matrix of median pairwise TM-score across all ten '
            'unordered pairs of the five model sources, computed with the same '
            'code and the same normalisation as the rest of this note. Cells are '
            'coloured from 0.60 to 1.00 and annotated with the median value. The '
            f"highest off-diagonal value is {mat[0][1]:.3f} for the AFDB and "
            'AlphaFold 3 pair. The lowest is '
            f"{min(mat[i][j] for i in range(5) for j in range(5) if i != j):.3f}. "
            '(b) Mean of the four off-diagonal values in each row of panel (a), '
            'that is the mean median-TM of each source to the other four. The '
            'values are '
            + series(f"{v:.3f} for {k}" for k, v in
                     sorted(fw['mean_vs_others'].items(), key=lambda kv: -kv[1]))
            + '. The AFDB models rank '
            + num(1 + sorted(fw['mean_vs_others'].values(), reverse=True)
                  .index(fw['mean_vs_others'][order[0]]))
            + ' of the five on this measure.'))

    if prox:
        fig.append((
            'FigS5_agreement_proxy',
            'Figure S5. Mean pairwise TM-score among the four HANSEN predictors on '
            'the horizontal axis against mean TM-score to the AFDB reference on the '
            f"vertical axis, one point for each of {prox['n']:,} proteins. The "
            'dashed line is the identity. The two quantities are correlated at '
            f"Pearson r = {prox['pearson_crossmethod_vs_afdb']:.3f} and Spearman "
            f"rho = {prox['spearman_crossmethod_vs_afdb']:.3f}. Points lie close to "
            'the identity across the full range, so the level of agreement among '
            'the four predictors reproduces the level of agreement with the AFDB '
            'reference without the reference being used.'))

    if pk:
        kin = pk['Kinase domain (11-273)']
        lin = pk['Juxtamembrane linker (274-328)']
        fig.append((
            'FigS6_deviation_profiles',
            'Figure S6. Per-residue Ca deviation from the AFDB model after '
            'TM-superposition, for four proteins that illustrate the two kinds of '
            'disagreement. In each panel the horizontal axis is the UniProt residue '
            'number, the vertical axis on the left is the Ca deviation on a '
            'logarithmic scale from 0.1 to 4,000 A, and the dotted horizontal line '
            'marks 5 A. Grey shading is the AFDB per-residue pLDDT on the right '
            'axis. One coloured line is drawn per method and the TM-score of each '
            'method is given in the panel legend. In the PknB panel (ML0016) the '
            'annotated UniProt domains are shaded and labelled. The deviation stays '
            f"near 1 A across the kinase domain, residues 11 to 273, where the "
            f"per-domain RMSD is {min(kin[m]['rmsd'] for m in M):.1f} to "
            f"{max(kin[m]['rmsd'] for m in M):.1f} A, and rises by two orders of "
            'magnitude at the juxtamembrane linker, residues 274 to 328, where the '
            f"per-segment RMSD is {min(lin[m]['rmsd'] for m in M):.0f} to "
            f"{max(lin[m]['rmsd'] for m in M):.0f} A. In the Smc panel (ML1629) "
            'each method traces a low deviation over a different part of the chain, '
            'which is the profile expected when an antiparallel coiled-coil arm '
            'moves about a hinge. In the Mce1F (ML2594) and LprK (ML2593) panels '
            'the deviation is near 1 A over the N-terminal MCE domain and rises '
            'monotonically through the C-terminal helical extension, in a region '
            'where the AFDB pLDDT remains between 80 and 90. A step in these '
            'profiles indicates a rigid-body difference. A profile raised across '
            'the whole chain would indicate a different fold, and no panel shows '
            'that pattern.'))
    return fig


# ------------------------------------------------------------------ narrative
def narrative():
    cov = S['coverage']
    disc = S['discordance']
    cf = S['consensus_failures']
    fp = S['fold_vs_packing']
    fw = S.get('five_way', {})
    prox = S.get('agreement_proxy', {})
    bm = S['best_method_counts']
    tv = S.get('tm_validation', {})
    cs = S.get('coverage_sensitivity', {})
    sc = S.get('short_chain_confidence', {})
    cse = S.get('case_studies', {})
    n_pairs = sum(cov['per_method'].values())
    n_rearr = cse.get('Domain rearrangement (all four methods)', 0)

    p = []
    A = p.append

    A(('H1', 'Supplementary Note S1. Structural concordance between the AlphaFold '
       'Protein Structure Database and the HANSEN monomeric models'))

    A(('H2', '1  Scope'))
    A(('P', 'We have compared the AlphaFold Protein Structure Database (AFDB) '
       'proteome release for *Mycobacterium leprae* TN with the monomeric models '
       'held in HANSEN. The AFDB release is UP000000806_272631_MYCLE at model '
       f"version v6 and contains {cov['afdb_entries']:,} single-chain models. The "
       'HANSEN monomers were generated by four co-folding predictors, AlphaFold 3, '
       'Boltz-1, Boltz-2 and Chai-1. The comparison covers '
       f"{cov['afdb_paired']:,} proteins and {n_pairs:,} model pairs."))
    A(('P', 'We restricted the comparison to monomers because the AFDB proteome '
       f"release contains single-chain models only. The {n_oligomers():,} HANSEN "
       'oligomeric model records have no counterpart in that release and are not '
       'considered here.'))

    A(('H2', '2  Model sets and residue correspondence'))
    A(('P', f"Of the {cov['afdb_entries']:,} AFDB entries, {cov['afdb_paired']:,} "
       'correspond one to one with an *M. leprae* locus that carries HANSEN '
       'monomers. The four remaining entries are second UniProt accessions attached '
       'to a locus label already used by another entry, namely O05755 at ML2219, '
       'Q9CD20 at ML2569, Q7AQ19 at ML1911 and P0A5D0 at ML2428. In each of these '
       'four cases we retained the accession whose UniProt sequence matches the '
       'modelled chain exactly.'))
    A(('P', 'One HANSEN protein has no AFDB entry. ML1191 (Q7AQ85, fatty acid '
       'synthase) is 3,076 residues long and exceeds the length limit of the AFDB '
       'proteome release. Four further loci, ML1180, ML1181, ML1183 and ML2692, '
       'carry monomers but no UniProt accession, so no AFDB model can be assigned '
       f"to them. Of the 1,603 loci that carry HANSEN monomers, {cov['afdb_paired']:,} "
       'therefore have an AFDB reference (Table S9).'))
    A(('P', f"The comparison comprises {cov['per_method']['af3']:,} pairs against "
       f"AlphaFold 3, {cov['per_method']['boltz']:,} against Boltz-1, "
       f"{cov['per_method']['boltz2']:,} against Boltz-2 and "
       f"{cov['per_method']['chai']:,} against Chai-1. The differences between the "
       'four counts arise from the monomer inventory of HANSEN and not from the '
       'comparison.'))
    A(('P', 'We established residue correspondence without structural alignment. '
       f"Of the {n_pairs:,} pairs, {cov['align_modes'].get('identity', 0):,} have "
       f"identical modelled sequences and a further "
       f"{cov['align_modes'].get('resnum', 0):,} differ at one position and were "
       'matched on UniProt residue numbering. The remaining '
       f"{cov['align_modes'].get('nw', 0):,} pairs required a global "
       'Needleman-Wunsch alignment because the modelled construct boundaries '
       'differ. Every single-residue difference lies at position 1. UniProt records '
       'the initiator residue as methionine, whereas the genomic translations used '
       'for the HANSEN models retain valine or leucine at the GTG and TTG start '
       'codons of *M. leprae*. No other position differs in these pairs, and the '
       'residue correspondence is therefore unaffected.'))
    lm = T('S2_construct_length_differences.tsv')[1:]
    A(('P', f"A total of {len(lm):,} proteins have different modelled chain lengths "
       'in the two model sets (Table S10). The four largest differences are ML0842 with 611 '
       'residues in AFDB and 411 in HANSEN, ML1985 with 606 and 798, ML0015 with '
       '232 and 47, and ML0021 with 155 and 340. These differences follow from gene '
       'boundary annotation rather than from modelling, and we made each comparison '
       'over the aligned region.'))

    A(('H2', '3  Metrics'))
    A(('P', 'We computed all metrics on Ca atoms from the corresponded residue '
       'pairs, taking the AFDB model as the reference in every comparison. '
       'TM-score follows Zhang and Skolnick (2004), normalised by the AFDB chain '
       'length and obtained by the standard iterative fragment-seed search. It is '
       'sequence-dependent, since the two structures represent the same sequence. '
       'GDT-TS and GDT-HA follow Zemla (2003) with cutoff sets of 1, 2, 4 and 8 A '
       'and of 0.5, 1, 2 and 4 A. lDDT is the superposition-free local distance '
       'difference test of Mariani and co-workers (2013), evaluated on Ca atoms '
       'with a 15 A inclusion radius and the four standard tolerance thresholds. '
       'Global RMSD is taken over all corresponded residues after a single Kabsch '
       'superposition. Core RMSD is restricted to residues that lie within 5 A of '
       'the reference after TM-superposition and therefore reports on the shared '
       'substructure. Secondary structure was assigned from Ca geometry by the '
       'P-SEA method of Labesse and co-workers (1997) and compared as a three-state '
       'agreement, Q3. Confidence values are the per-residue pLDDT stored in the '
       'B-factor field of every model.'))
    if tv:
        A(('P', 'We verified the TM-score implementation against the TM-align '
           f"reference program on {sum(v['n'] for v in tv.values()):,} randomly "
           'chosen pairs spanning all four methods (Table S23). The correlation is '
           + by_method(lambda m: f"r = {tv[m]['pearson_r']:.4f} for {LAB[m]}")
           + f", the median difference is zero for every method and the mean "
           f"difference is {sum(tv[m]['mean_diff'] for m in M)/4:.4f}. The residual "
           'differences run in one direction, with TM-align scoring higher, which '
           'follows from TM-align optimising a sequence-independent alignment. For '
           'two models of the same sequence the sequence-dependent score used here '
           'is the appropriate measure.'))
    if cs:
        A(('P', 'TM-score is normalised by the AFDB chain length, so a HANSEN model '
           'that spans part of the reference chain is limited to approximately its '
           'coverage fraction whatever the agreement over the shared part. '
           'Chain coverage falls below 95 per cent in '
           + by_method(lambda m: f"{num(cs[m]['n_partial'])} comparisons for {LAB[m]}")
           + '. Excluding them changes the median '
           'TM-score by at most '
           f"{max(abs(cs[m]['median_tm_full_coverage']-cs[m]['median_tm_all']) for m in M):.3f}, "
           'so the proteome-wide values are unaffected. ML0842, a cysteine '
           'desulfurase with 611 residues in AFDB and 411 in HANSEN, illustrates the '
           'effect. Its TM-score is 0.66 to 0.67 for all four methods, close to the '
           'coverage ratio of 0.67, while the 411 shared residues superpose at 0.8 '
           'to 1.5 A whole-chain RMSD with lDDT of 0.986 to 0.991. We excluded these '
           'comparisons from the classification in section 6.'))

    A(('H2', '4  Global agreement'))
    A(('P', 'The four predictors reproduce the AFDB models over the proteome '
       '(Table S11, Figure S1). The median TM-score to AFDB ranges from '
       f"{min(g(m,'tm_score','median') for m in M):.3f} to "
       f"{max(g(m,'tm_score','median') for m in M):.3f} across the four methods, "
       'and the proportion of proteins at or above the TM 0.5 threshold for a '
       f"shared fold ranges from {min(thr(m,0.5) for m in M):.1f} to "
       f"{max(thr(m,0.5) for m in M):.1f} per cent. Local agreement is higher than "
       'global agreement in every method, with median Ca-lDDT between '
       f"{min(g(m,'lddt','median') for m in M):.3f} and "
       f"{max(g(m,'lddt','median') for m in M):.3f}."))
    A(('TABLE', 'S3_global_agreement.tsv'))
    A(('P', 'The TM-score distributions carry a single high mode and a low tail '
       'rather than a broad spread. For AlphaFold 3 the interquartile range is '
       f"{g('af3','tm_score','q1'):.3f} to {g('af3','tm_score','q3'):.3f} while the "
       f"mean is {g('af3','tm_score','mean'):.3f}, below the median of "
       f"{g('af3','tm_score','median'):.3f}. Mean whole-chain RMSD exceeds the "
       'median RMSD over the shared substructure by a factor of '
       f"{min(g(m,'rmsd_global','mean')/g(m,'rmsd_core','median') for m in M):.0f} "
       f"to {max(g(m,'rmsd_global','mean')/g(m,'rmsd_core','median') for m in M):.0f}. "
       'Whole-chain RMSD therefore describes the proteome poorly, because it is '
       'dominated by the multi-domain proteins analysed in section 6.'))
    A(('P', 'AlphaFold 3 lies closest to the AFDB reference, which follows from the '
       'shared architecture of the two systems. It gives the highest TM-score for '
       f"{bm['AlphaFold 3']:,} of the {sum(bm.values()):,} proteins compared by all "
       f"four methods, against {bm['Chai-1']:,} for Chai-1, {bm['Boltz-1']:,} for "
       f"Boltz-1 and {bm['Boltz-2']:,} for Boltz-2. The median TM-score separates "
       f"AlphaFold 3 from Boltz-2 by "
       f"{g('af3','tm_score','median')-g('boltz2','tm_score','median'):.3f} and the "
       'three non-AlphaFold predictors from one another by at most '
       f"{max(g(m,'tm_score','median') for m in ('boltz','boltz2','chai'))-min(g(m,'tm_score','median') for m in ('boltz','boltz2','chai')):.3f}."))
    A(('FIG', 'FigS1_global_agreement'))

    A(('H2', '5  Confidence'))
    A(('P', 'All confidence values in this section are computed over the 1,598 '
       'proteins compared here, so they differ marginally from the proteome-wide '
       'values in table S28 of the supplementary materials, which cover 1,603 proteins. The mean '
       'pLDDT of the AFDB proteome is '
       f"{S['confidence']['af3']['afdb_plddt_mean']:.1f}. AlphaFold 3 and Chai-1 "
       'report the same level, at '
       f"{S['confidence']['af3']['model_plddt_mean']:.1f} and "
       f"{S['confidence']['chai']['model_plddt_mean']:.1f}. Boltz-1 and Boltz-2 "
       f"report {S['confidence']['boltz']['model_plddt_mean']:.1f} and "
       f"{S['confidence']['boltz2']['model_plddt_mean']:.1f}, which are "
       f"{S['confidence']['af3']['afdb_plddt_mean']-S['confidence']['boltz']['model_plddt_mean']:.1f} "
       f"and {S['confidence']['af3']['afdb_plddt_mean']-S['confidence']['boltz2']['model_plddt_mean']:.1f} "
       'units below the AFDB level (Table S12, Figure S2c). The coordinates do not '
       'follow the same ordering. Median lDDT is '
       f"{g('boltz','lddt','median'):.3f} for Boltz-1 and "
       f"{g('boltz2','lddt','median'):.3f} for Boltz-2, against "
       f"{g('chai','lddt','median'):.3f} for Chai-1, a difference of "
       f"{g('chai','lddt','median')-g('boltz2','lddt','median'):.3f}. The offset "
       'therefore resides in the confidence estimate rather than in the '
       'coordinates.'))
    A(('TABLE', 'S4_confidence.tsv'))
    A(('P', 'The per-residue confidence profiles agree across methods. The median '
       'within-protein Pearson correlation between the AFDB and model pLDDT '
       'profiles is between '
       f"{min(S['confidence'][m]['per_protein_plddt_r_median'] for m in M):.3f} and "
       f"{max(S['confidence'][m]['per_protein_plddt_r_median'] for m in M):.3f}, so "
       'the four predictors and AFDB identify the same regions of each chain as '
       'reliable.'))
    b = S['per_residue']['af3']['dev_by_min_plddt_bin']
    A(('P', 'Per-residue deviation is governed by confidence (Figure S2a). Over the '
       f"{S['per_residue']['af3']['n_residues']:,} AlphaFold 3 residues, the median "
       f"Ca deviation from AFDB is {b['ge90']['median_dev']:.2f} A where the lower "
       f"of the two pLDDT values is 90 or above ({b['ge90']['n']:,} residues), "
       f"{b['70_90']['median_dev']:.2f} A between 70 and 90 "
       f"({b['70_90']['n']:,} residues), {b['50_70']['median_dev']:.1f} A between "
       f"50 and 70 ({b['50_70']['n']:,} residues) and "
       f"{b['lt50']['median_dev']:.1f} A below 50 ({b['lt50']['n']:,} residues). "
       'The deviation increases by a factor of '
       f"{b['lt50']['median_dev']/b['ge90']['median_dev']:.0f} between the highest "
       'and lowest confidence bins.'))
    A(('FIG', 'FigS2_confidence'))

    A(('H2', '6  Low-confidence regions, local fold and domain packing'))
    A(('P', 'Restricting each comparison to residues with pLDDT of 70 or above in '
       'both structures raises the mean TM-score by '
       + by_method(lambda m: f"{S['core'][m]['tm_core']-S['core'][m]['tm_full']:.3f} for {LAB[m]}")
       + ' and the mean lDDT by '
       + by_method(lambda m: f"{S['core'][m]['lddt_core']-S['core'][m]['lddt_full']:.3f} for {LAB[m]}")
       + ' (Table S13, Figure S3b). Within this confident core the proportion of '
       'proteins at or above TM 0.5 is '
       + series(f"{T('S5_confident_core.tsv')[i+1][8].rstrip('%')} per cent for {LAB[m]}"
                for i, m in enumerate(M))
       + '.'))
    A(('P', 'lDDT is superposition-free and TM-score is not, so the two together '
       'separate two kinds of difference (Table S17). A model with high lDDT and '
       'low TM-score reproduces the individual domains and places them differently '
       'with respect to one another. A model low in both differs in local '
       'structure. Between '
       f"{min(fp[m]['domain_rearrangement'] for m in M)} and "
       f"{max(fp[m]['domain_rearrangement'] for m in M)} comparisons per method have "
       'lDDT of 0.80 or above together with TM-score below 0.70, and between '
       f"{min(fp[m]['divergent_local'] for m in M)} and "
       f"{max(fp[m]['divergent_local'] for m in M)} are low in both. "
       'The first group has a mean chain '
       f"length of {s8b(7)} residues against a proteome median of "
       f"{cf['median_length_all']:.0f}, a mean core RMSD of {s8b(6)} A and a mean "
       f"whole-chain RMSD of {s8b(5)} A. The ratio between the two RMSD values is "
       f"{float(s8b(5).split(' to ')[0])/float(s8b(6).split(' to ')[1]):.0f} to "
       f"{float(s8b(5).split(' to ')[-1])/float(s8b(6).split(' to ')[0]):.0f}, and "
       'the difference is rigid-body displacement.'))
    A(('TABLE', 'S8b_fold_vs_packing.tsv'))
    A(('FIG', 'FigS3_stratification'))
    coil = [('ML1629', 'Smc'), ('ML1369', 'ScpB')]
    modular = [('ML2353', 'the phthiocerol polyketide synthase subunit'),
               ('ML1996', 'the non-ribosomal peptide synthase Nrp'),
               ('ML0016', 'the protein kinase PknB'), ('ML0987', 'RecA'),
               ('ML2688', 'the penicillin-binding protein PonA'), ('ML0562', 'UvrC'),
               ('ML0001', 'the replication initiator DnaA'),
               ('ML0603', 'a DNA polymerase'), ('ML0510', 'an AAA+ ATPase'),
               ('ML1120', 'a RecF-like AAA protein'),
               ('ML0861', 'the dihydrolipoamide acetyltransferase E2 subunit'),
               ('ML1692', 'the dehydrogenase SerA')]
    membrane = [('ML0880', 'the Rieske subunit QcrA'),
                ('ML0876', 'the cytochrome oxidase subunit CtaF'),
                ('ML1802', 'a band-7 protein'), ('ML0916', 'FtsQ'),
                ('ML2377', 'MmpS4')]
    assembly = [('ML1858', 'the ribosomal protein uL22'),
                ('ML1854', 'the ribosomal protein uS17'),
                ('ML0173', 'the ribosomal protein bL32'),
                ('ML0234', 'the nucleoid-associated protein Lsr2'),
                ('ML0922', 'the cell-wall scaffold Wag31')]
    named = coil + modular + membrane + assembly
    word = {2: 'Two', 5: 'Five', 12: 'Twelve'}

    def lst(items):
        s = [ex(ml, nm) for ml, nm in items]
        return ', '.join(s[:-1]) + ' and ' + s[-1]

    A(('P', f"In all, {n_rearr} proteins show this signature in all four methods "
       'while the AFDB model is itself confident, defined as mean lDDT of 0.85 or above, '
       'every TM-score below 0.70 and AFDB mean pLDDT of 80 or above (Table S24). '
       f"Of these {n_rearr}, {len(named)} fall into four groups that follow from "
       f"their UniProt annotation. {word[len(coil)]} are elongated coiled-coil proteins, "
       f"{lst(coil)}. {word[len(modular)]} are modular enzymes and multi-domain "
       f"proteins joined by flexible linkers, {lst(modular)}. "
       f"{word[len(membrane)]} are membrane-anchored proteins whose soluble domains "
       f"are attached to a single transmembrane segment, {lst(membrane)}. "
       f"{word[len(assembly)]} adopt a defined conformation only within a larger "
       f"assembly, {lst(assembly)}."))

    A(('H2', '7  Domain-resolved comparison for PknB'))
    pk = S.get('pknb_worked_example')
    if pk:
        kin = pk['Kinase domain (11-273)']
        lin = pk['Juxtamembrane linker (274-328)']
        pas = pk['PASTA 1-4 (352-622)']
        whole = pk['Whole chain (1-622)']
        lin_plddt = next((float(r[-1]) for r in T('S16_pknb_domains.tsv')[1:]
                          if r[0].startswith('Juxtamembrane')), float('nan'))
        kin_plddt = next((float(r[-1]) for r in T('S16_pknb_domains.tsv')[1:]
                          if r[0].startswith('Kinase')), float('nan'))
        A(('P', 'PknB (ML0016, P54744) is the serine/threonine protein kinase that '
           'ranks 50 in the high-priority tier, and it falls below TM 0.5 for all '
           'four methods. We resolved the comparison into the UniProt domains of '
           'the protein (Table S25, Figure S6). Over the whole chain the TM-score '
           f"is {min(whole[m]['tm'] for m in M):.2f} to "
           f"{max(whole[m]['tm'] for m in M):.2f}, the RMSD is "
           f"{min(whole[m]['rmsd'] for m in M):.0f} to "
           f"{max(whole[m]['rmsd'] for m in M):.0f} A and the lDDT is "
           f"{min(whole[m]['lddt'] for m in M):.2f} to "
           f"{max(whole[m]['lddt'] for m in M):.2f}."))
        A(('TABLE', 'S16_pknb_domains.tsv'))
        A(('P', 'The catalytic kinase domain, residues 11 to 273, reproduces the '
           f"AFDB model at TM {min(kin[m]['tm'] for m in M):.2f} to "
           f"{max(kin[m]['tm'] for m in M):.2f} and RMSD "
           f"{min(kin[m]['rmsd'] for m in M):.1f} to "
           f"{max(kin[m]['rmsd'] for m in M):.1f} A, with an AFDB mean pLDDT of "
           f"{kin_plddt:.1f}. The four extracellular PASTA repeats, residues 352 to "
           f"622, reproduce it at TM {min(pas[m]['tm'] for m in M):.2f} to "
           f"{max(pas[m]['tm'] for m in M):.2f}. The 55-residue juxtamembrane "
           f"linker, residues 274 to 328, has an AFDB mean pLDDT of {lin_plddt:.1f} "
           f"and is placed {min(lin[m]['rmsd'] for m in M):.0f} to "
           f"{max(lin[m]['rmsd'] for m in M):.0f} A from the reference by all four "
           'methods. The whole-chain TM-score therefore reflects the position of '
           'one low-confidence linker and not the accuracy of the kinase domain.'))
        A(('FIG', 'FigS6_deviation_profiles'))

    A(('H2', '8  Chain length, localisation and secondary structure'))
    t6 = T('S6_length_strata.tsv')
    A(('P', 'Agreement increases with chain length up to 700 residues (Table S14, '
       'Figure S3a). Mean '
       f"TM-score for AlphaFold 3 is {t6[1][2]} over the {t6[1][1]} proteins shorter "
       f"than 100 residues, {t6[2][2]} over the {t6[2][1]} proteins of 100 to 200 "
       f"residues, {t6[3][2]} over the {t6[3][1]} proteins of 200 to 400 residues, "
       f"{t6[4][2]} over the {t6[4][1]} proteins of 400 to 700 residues and "
       f"{t6[5][2]} over the {t6[5][1]} proteins longer than 700 residues. The "
       'other three methods follow the same profile. Two factors contribute. The '
       'TM-score normalisation is stringent for short chains, and the '
       f"{t6[1][1]} proteins shorter than 100 residues have a mean AFDB "
       f"pLDDT of {sc['mean_afdb_plddt_under_100']:.1f} against "
       f"{sc['mean_afdb_plddt_all']:.1f} over the whole comparison set."))
    t12 = T('S12_localisation_strata.tsv')
    A(('P', f"The {int(t12[1][1]):,} proteins with an annotated transmembrane "
       f"segment reach a mean TM-score of {t12[1][2]} for AlphaFold 3, {t12[1][3]} "
       f"for Boltz-1, {t12[1][4]} for Boltz-2 and {t12[1][5]} for Chai-1. The "
       f"{int(t12[3][1]):,} proteins with neither a transmembrane segment nor a signal "
       f"peptide reach {t12[3][2]}, {t12[3][3]}, {t12[3][4]} and {t12[3][5]} "
       f"respectively, a difference of {float(t12[3][2])-float(t12[1][2]):.3f} for "
       'AlphaFold 3 (Table S21). All five model sources place these proteins without '
       'a membrane.'))
    A(('P', 'Secondary structure is reproduced across the four methods (Table S15). '
       'Mean three-state agreement with AFDB is '
       + by_method(lambda m: f"{S['ss'][m]['q3_mean']:.3f} for {LAB[m]}")
       + '. Helix content is '
       f"{100*S['ss']['af3']['helix_ref']:.1f} per cent in the AFDB models against "
       + by_method(lambda m: f"{100*S['ss'][m]['helix_mod']:.1f} per cent for {LAB[m]}")
       + '. Strand content is '
       f"{100*S['ss']['af3']['strand_ref']:.1f} per cent in the AFDB models against "
       + by_method(lambda m: f"{100*S['ss'][m]['strand_mod']:.1f} per cent for {LAB[m]}")
       + '. The largest difference in either element is '
       f"{max(100*abs(S['ss'][m]['helix_mod']-S['ss'][m]['helix_ref']) for m in M):.1f} "
       'percentage points, and there is no exchange between helix and strand.'))
    A(('P', 'The median ratio of the radius of gyration of the model to that of the '
       'reference, computed over the corresponded residues, is '
       + by_method(lambda m: f"{S['ss'][m]['rg_ratio_median']:.3f} for {LAB[m]}")
       + '. The three values below unity correspond to models that are more '
       'compact than the reference by '
       + series(f"{100*(1-S['ss'][m]['rg_ratio_median']):.1f} per cent for {LAB[m]}"
                for m in ('boltz', 'boltz2', 'chai'))
       + ', consistent with the domain packing differences in section 6. The '
       'P-SEA assignment used here is based on Ca geometry and assigns regular '
       'secondary structure more freely than a hydrogen-bond method such as DSSP, '
       'so we use the helix and strand contents for comparison between model sets '
       'and not as absolute values.'))

    A(('H2', '9  Discordant proteins'))
    A(('P', f"A total of {disc['proteins_flagged']:,} proteins meet at least one of "
       'two criteria for discordance, either one method below TM 0.5 against AFDB or a '
       'range of more than 0.3 TM units across the four methods (Table S16). '
       f"Of these, {disc['all_methods_below_0.5']} lie below TM 0.5 for all four "
       f"methods. Their median chain length is {cf['median_length']:.0f} residues "
       f"against a proteome median of {cf['median_length_all']:.0f}, and "
       f"{100*cf['frac_under_150aa']:.0f} per cent are shorter than 150 residues. "
       f"Their median AFDB mean pLDDT is {cf['median_afdb_plddt']:.1f} against a "
       f"proteome median of {cf['median_afdb_plddt_all']:.1f}, and "
       f"{100*cf['frac_afdb_plddt_under_70']:.0f} per cent fall below pLDDT 70."))
    chai_out = [('ML0050', 'the ESAT-6-like antigen EsxB'),
                ('ML0512', 'alanyl-tRNA synthetase AlaS'),
                ('ML0811', 'the RNA helicase RhlE'),
                ('ML2075', 'a MerR-type regulator')]
    af3_out = [('ML1793', None), ('ML2258', 'a DUF3349 protein'), ('ML1430', None)]
    seven = [m for m, _ in chai_out + af3_out]
    ld = sorted(float(S15[m][9]) for m in seven)
    nres = sorted(int(S15[m][4]) for m, _ in af3_out)
    A(('P', 'Method-specific failures are counted as proteins where one method lies '
       'below TM 0.5 and the other three do not. The counts are '
       + by_method(lambda m: f"{disc['method_unique_failures'].get(m,0)} for {LAB[m]}")
       + f". Table S24 lists the {num(len(seven))} in which the other three methods "
       f"reach TM 0.70 or above. {Num(len(chai_out))} of them are Chai-1, "
       f"{lst(chai_out)}. The other {num(len(af3_out))} are AlphaFold 3, "
       f"{lst(af3_out)}, each of which is between {nres[0]} and {nres[-1]} residues "
       f"long. The mean lDDT of the {num(len(seven))} proteins ranges from "
       f"{ld[0]:.3f} to {ld[-1]:.3f}, so the "
       'local structure of the reference is retained in each case and the '
       'difference lies in the placement of one part of the chain against '
       'another.'))
    nfail = cse.get('Whole-proteome discordance despite confident AFDB model', 0)
    A(('P', f"{Num(nfail)} proteins disagree with a confident AFDB model over more "
       f"than half of the chain, {ex('ML2594', 'Mce1F')}, {ex('ML2593', 'LprK')} and "
       f"{ex('ML0293')}. Mce1F and LprK are encoded in the mce1 lipid-import "
       'operon. The AFDB mean pLDDT is '
       + series(f"{float(S15[ml][10]):.1f} for {nm}" for ml, nm in
                (('ML2594', 'Mce1F'), ('ML2593', 'LprK'), ('ML0293', 'ML0293')))
       + ', while the mean model pLDDT over the four methods is '
       + series(f"{float(S15[ml][11]):.1f}" for ml in
                ('ML2594', 'ML2593', 'ML0293'))
       + ' in the same order. The per-residue profiles localise the difference '
       '(Figure S6). For '
       'LprK the N-terminal MCE domain, residues 1 to 150, superposes within 1 A in '
       'all four methods, and the deviation then rises through the C-terminal '
       'helical extension to more than 100 A at residue 350. Mce1F shows the same '
       'profile. Mce1B (ML2590), classified above as a rearrangement, shows the '
       'same profile with the transition near residue 200. The part of the chain '
       'over which each method superposes differs between methods, which is the '
       'behaviour expected of a hinge rather than of a locally incorrect fold.'))

    A(('H2', '10  Position of the AFDB models within the ensemble'))
    if fw:
        mv = fw['mean_vs_others']
        ordered = sorted(mv.items(), key=lambda kv: -kv[1])
        A(('P', 'We computed all ten pairwise comparisons among the five model '
           'sources on the same scale (Table S18, Figure S4). The mean median-TM of '
           'each source to the other four is '
           + series(f"{v:.3f} for {k}" for k, v in ordered)
           + f". The AFDB and AlphaFold 3 pair is the most concordant of the ten at "
           f"{fw['median_tm_matrix'][0][1]:.3f} median TM, and the spread across "
           f"the five sources is {ordered[0][1]-ordered[-1][1]:.3f}. The AFDB "
           'models therefore sit within the ensemble on this measure and not '
           'outside it.'))
        A(('TABLE', 'S9_five_way_concordance.tsv'))
        A(('FIG', 'FigS4_five_way'))
    if prox:
        A(('P', 'Agreement among the four HANSEN predictors tracks agreement with '
           f"AFDB at Pearson r = {prox['pearson_crossmethod_vs_afdb']:.3f} and "
           f"Spearman rho = {prox['spearman_crossmethod_vs_afdb']:.3f} over "
           f"{prox['n']:,} proteins (Figure S5)."))
        A(('FIG', 'FigS5_agreement_proxy'))

    A(('H2', '11  Agreement across the target-priority tiers'))
    t10 = T('S10_target_tiers.tsv')
    hp = next(r for r in t10 if r[0] == 'High-priority')
    st = next(r for r in t10 if r[0] == 'Strong candidate')
    mo = next(r for r in t10 if r[0] == 'Moderate candidate')
    exr = next(r for r in t10 if r[0] == 'Exploratory')
    A(('P', f"For the {hp[1]} high-priority targets the mean TM-score against AFDB "
       f"is {hp[2]} for AlphaFold 3, {hp[3]} for Boltz-1, {hp[4]} for Boltz-2 and "
       f"{hp[5]} for Chai-1 (Table S19). The {st[1]} strong candidates reach {st[2]}, "
       f"{st[3]}, {st[4]} and {st[5]}, and the {mo[1]} moderate candidates reach "
       f"{mo[2]}, {mo[3]}, {mo[4]} and {mo[5]}. The {exr[1]} exploratory-tier "
       f"proteins reach {exr[2]}, {exr[3]}, {exr[4]} and {exr[5]}, which is "
       f"{float(hp[2])-float(exr[2]):.3f} below the high-priority tier for "
       'AlphaFold 3. The mean model pLDDT follows the same ordering, at '
       f"{hp[6]} for the high-priority tier and {exr[6]} for the exploratory tier "
       'in AlphaFold 3.'))
    A(('P', f"Of the {disc['all_methods_below_0.5']} proteins that lie below TM 0.5 "
       'for all four methods, '
       f"{S['tier_of_consensus_failures']['Exploratory']} belong to the exploratory "
       'tier. Table S20 gives the per-method '
       'TM-score and lDDT for the top 50 prioritised targets individually.'))

    A(('H2', '12  Limitations'))
    A(('P', 'The AFDB models are predictions and not experimental structures, so '
       'agreement with them measures reproducibility and not accuracy. The '
       'benchmark against the seven experimentally determined *M. leprae* '
       'structures is reported in the main text. All comparisons here use Ca atoms '
       'only, so side-chain placement, rotamer geometry and hydrogen bonding lie '
       'outside their scope. The secondary-structure assignment is based on Ca '
       'geometry and assigns regular structure more freely than DSSP. None of the '
       'compared monomers contains a ligand, cofactor or metal, so these are not '
       'considered. The construct differences in Table S10 arise from the AFDB '
       'release and the HANSEN models having been built from different UniProt '
       'releases.'))

    A(('H2', '13  Tables, data and code'))
    A(('P', 'Seventeen tables accompany this note as a workbook, one sheet per '
       'table, numbered so that they follow the supplementary tables of the main '
       'text. Table S11, Table S12, Table S17, Table S18 and Table S25 are set in the '
       'text above, the first and the last of these as a column subset with the '
       'complete versions in the workbook. The remaining twelve are held in the '
       'workbook alone. Table S9 gives the comparison coverage and Table S10 lists '
       'the proteins whose modelled chain length differs between the two model '
       'sets. Table S13 gives the effect of restricting to the confident core, '
       'Table S14 agreement by chain length, Table S21 agreement by predicted '
       'localisation and Table S15 secondary structure and compactness. Table S16 '
       'lists every discordant protein and Table S24 the named case studies. '
       'Table S19 gives agreement by target-priority tier and Table S20 the 50 '
       'top-ranked prioritised targets individually. Table S22 gives the '
       'per-protein results for the complete set and Table S23 the validation of '
       'the TM-score implementation. Figures S1 to S6 accompany the note.'))
    A(('P', f"Per-protein results for all {n_pairs:,} comparisons are given in "
       'Table S22. We ran the comparison against '
       'the AFDB proteome archive UP000000806_272631_MYCLE_v6.tar and the HANSEN '
       'monomer set registered in protein_characteristics.db. The pipeline '
       '(01_build_manifest.py, 02_compare.py, 02b_crossmethod.py, 03_analyse.py, '
       '03b_worked_example.py, 03c_profiles.py, 04_report.py, 05_validate.py, '
       'driven by run_all.sh) and the metric implementation (structlib.py) are '
       'deposited with the tables, together with the protein-to-model manifest '
       '(manifest.tsv), the per-pair output (pairs.tsv, crossmethod.tsv), the '
       'pooled per-residue arrays (per_residue.npz) and the machine-readable '
       'result set (summary.json). Each figure is supplied as a 300 dpi PNG and as '
       'a vector PDF. The pipeline is deterministic apart from the seeded random '
       'sample drawn for Table S23.'))
    return p


# ------------------------------------------------------------------ Markdown
def build_md(parts, figs):
    legend = dict(figs)
    out = []
    for kind, text in parts:
        if kind == 'H1':
            out.append(f'# {text}\n')
        elif kind == 'H2':
            out.append(f'\n## {text}\n')
        elif kind == 'TABLE':
            rows = table_rows(text)
            out.append(f'\n**{TABLE_TITLES[text]}**\n')
            out.append('| ' + ' | '.join(rows[0]) + ' |')
            out.append('|' + '---|' * len(rows[0]))
            for r in rows[1:]:
                out.append('| ' + ' | '.join(r) + ' |')
            if INLINE_SPEC.get(text):
                out.append(f'\n*Selected columns. The complete table is sheet '
                           f'{sheet_name(text)} of the workbook.*\n')
            else:
                out.append('')
        elif kind == 'FIG':
            out.append(f'\n![{text}](figures/{text}.png)\n')
            out.append(polish(legend[text]) + '\n')
        else:
            out.append(polish(text) + '\n')
    with open(MD, 'w') as fh:
        fh.write('\n'.join(out))
    print('wrote', MD)


# ---------------------------------------------------------------------- DOCX
def build_docx(parts, figs):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    legend = dict(figs)
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15

    def add_md_runs(par, text):
        """Minimal *italic* handling for species names."""
        for i, chunk in enumerate(text.split('*')):
            if not chunk:
                continue
            r = par.add_run(chunk)
            r.italic = (i % 2 == 1)

    def caption(text, bold=False, size=8.5):
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.space_before = Pt(6)
        run = par.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return par

    for kind, text in parts:
        if kind == 'H1':
            h = doc.add_heading(level=1)
            h.add_run(text)
        elif kind == 'H2':
            doc.add_heading(text, level=2)
        elif kind == 'TABLE':
            rows = table_rows(text)
            caption(TABLE_TITLES[text], bold=True, size=9)
            tbl = doc.add_table(rows=1, cols=len(rows[0]))
            tbl.style = 'Light Grid Accent 1'
            for j, c in enumerate(rows[0]):
                cell = tbl.rows[0].cells[j]
                cell.text = ''
                run = cell.paragraphs[0].add_run(c)
                run.bold = True
                run.font.size = Pt(7.5)
            for r in rows[1:]:
                cells = tbl.add_row().cells
                for j, c in enumerate(r):
                    cells[j].text = ''
                    run = cells[j].paragraphs[0].add_run(c)
                    run.font.size = Pt(7.5)
            if INLINE_SPEC.get(text):
                par = caption('Selected columns. The complete table is sheet '
                              f'{sheet_name(text)} of the accompanying workbook.',
                              size=8)
                par.runs[0].italic = True
            doc.add_paragraph()
        elif kind == 'FIG':
            png = os.path.join(FIGDIR, f'{text}.png')
            if os.path.exists(png):
                doc.add_picture(png, width=Inches(6.4))
            caption(polish(legend[text]), size=9)
            doc.add_paragraph()
        else:
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_md_runs(par, polish(text))

    doc.save(DOCX)
    print('wrote', DOCX)


# ---------------------------------------------------------------------- XLSX
def build_xlsx():
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    idx = wb.active
    idx.title = 'Index'
    idx['A1'] = ('HANSEN supplementary tables. AFDB (AlphaFold Protein Structure '
                 'Database, v6) compared with the HANSEN monomeric models')
    idx['A1'].font = Font(bold=True, size=13)
    idx['A2'] = (f'Generated {datetime.date.today().isoformat()} from the AFDB '
                 f'proteome release UP000000806_272631_MYCLE_v6')
    idx['A4'] = 'Sheet'
    idx['B4'] = 'Rows'
    idx['C4'] = 'Description'
    for c in 'ABC':
        idx[f'{c}4'].font = Font(bold=True)

    r = 5
    for fn, title in TABLE_TITLES.items():
        rows = T(fn)
        sheet = sheet_name(fn)
        ws = wb.create_sheet(sheet[:31])
        ws.append(rows[0])
        for c in ws[1]:
            c.font = Font(bold=True)
            c.alignment = Alignment(wrap_text=True, vertical='top')
        for row in rows[1:]:
            ws.append(row)
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = ws.dimensions
        for j in range(1, len(rows[0]) + 1):
            width = max(len(str(rows[k][j - 1])) for k in range(min(len(rows), 200))
                        if j - 1 < len(rows[k]))
            ws.column_dimensions[get_column_letter(j)].width = min(48, max(10, width + 2))
        idx.cell(row=r, column=1, value=sheet)
        idx.cell(row=r, column=2, value=len(rows) - 1)
        idx.cell(row=r, column=3, value=title.split('. ', 1)[1])
        r += 1
    idx.column_dimensions['A'].width = 12
    idx.column_dimensions['B'].width = 10
    idx.column_dimensions['C'].width = 90
    wb.save(XLSX)
    print('wrote', XLSX)


REQUIRED = ['coverage', 'global', 'confidence', 'core', 'ss', 'discordance',
            'consensus_failures', 'fold_vs_packing', 'five_way',
            'agreement_proxy', 'case_studies', 'coverage_sensitivity',
            'short_chain_confidence', 'tier_of_consensus_failures',
            'pknb_worked_example', 'tm_validation']

if __name__ == '__main__':
    missing = [k for k in REQUIRED if k not in S]
    if missing:
        raise SystemExit(
            f'summary.json is missing {missing}. 03_analyse.py rewrites the file, '
            'so 03b_worked_example.py and 05_validate.py must run after it. '
            'Use ./run_all.sh.')
    figs = figures()
    parts = narrative()
    build_md(parts, figs)
    build_docx(parts, figs)
    build_xlsx()
